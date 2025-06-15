import time

from docx import Document
from docx.shared import Pt

def set_font_with_auto_shrink(run, text, max_length=36):
    try:
        run.text = text
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8) if len(text) > max_length else Pt(9)
    except Exception as e:
        print(f"[ERROR] Ошибка в set_font_with_auto_shrink: {e}")


def replace_in_paragraphs(paragraphs, values):
    try:
        for para in paragraphs:
            for run in para.runs:
                original_text = run.text
                new_text = original_text
                for key, val in values.items():
                    new_text = new_text.replace(key, str(val))
                if new_text != original_text:
                    set_font_with_auto_shrink(run, new_text)
    except Exception as e:
        print(f"[ERROR] Ошибка в replace_in_paragraphs: {e}")

def replace_in_tables(tables, values):
    try:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs, values)
                    if cell.tables:
                        replace_in_tables(cell.tables, values)
    except Exception as e:
        print(f"[ERROR] Ошибка в replace_in_tables: {e}")

def generate_document(values: dict):
    try:
        doc = Document("assets/Заявление_СПП.docx")
        replace_in_paragraphs(doc.paragraphs, values)
        replace_in_tables(doc.tables, values)
        output_path = f"output_{int(time.time())}.docx"
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"[ERROR] Ошибка в replace_in_tables: {e}")
        return None
